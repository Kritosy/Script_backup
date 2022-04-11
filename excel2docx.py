import xlrd
import docx
'''
    excel格式
    +-----+-----+------+--------------------+--------+---------+-----+-----+
    | 学校 | 姓名 | 性别 | 身份证号             | 父母姓名 | 手机号码 | 年级 | 班级 |
    +-----+-----+------+--------------------+--------+---------+-----+-----+
    | 北京 | 张三 | 男   | 123456789012345678 | 张四    | 123XXXX | 高一 |12班 |
    +-----+-----+------+--------------------+--------+---------+-----+-----+
'''
database = xlrd.open_workbook('demo.xls')
table = database.sheet_by_name('Sheet1')
document = docx.Document()

#设置读取文字居中
def text_center_8(text):
    length = len(text)
    if length == 8:
        return text
    elif length == 7:
        return " " + text + " "
    elif length == 6:
        return "  " + text + "  "
    elif length == 5:
        return "   " + text + "   "
    elif length == 4:
        return "    " + text + "    "
    elif length == 3:
        return "     " + text + "     "
    elif length == 2:
        return "      " + text + "      "
    elif length == 1:
        return "       " + text + "       "

def text_center_4(text):
    length = len(text)
    if length == 4:
        return text
    elif length == 3:
        return " " + text + " "
    elif length == 2:
        return "  " + text + "  "
    elif length == 1:
        return "   " + text + "   "

# 设置页边距
for section in document.sections:
    section.top_margin = docx.shared.Cm(1.5)
    section.bottom_margin = docx.shared.Cm(2)
    section.left_margin = docx.shared.Cm(2)
    section.right_margin = docx.shared.Cm(2)

rows = table.nrows


for i in range(rows):
    school = table.cell_value(i, 0)
    school_center = text_center_8(school)
    name = table.cell_value(i, 1)
    sex = table.cell_value(i, 2)
    ID_num = table.cell_value(i, 3)
    ParentName = table.cell_value(i, 4)
    phone = table.cell_value(i, 5)

    Id_num = list(ID_num)
    birthday = "".join(Id_num[6:14])

    para = document.add_paragraph()
    para.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.SINGLE
    para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    para_run = para.add_run("投保回执单")
    para_run.font.size = docx.shared.Pt(12)
    para_run.font.name = "仿宋"
    para_run.font.bold = True
    para_run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)

    para1 = document.add_paragraph()
    para1.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.SINGLE
    p_run = para1.add_run()
    p_run.font.size = docx.shared.Pt(9)
    p_run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    p1 = para1.add_run("    投保声明：我已阅知本保险条款、责任免除、理赔须知等内容，自愿为")
    p1.font.size = docx.shared.Pt(9)

    p2 = para1.add_run(school_center)
    p2.font.underline = True
    p2.font.size = docx.shared.Pt(9)

    p3 = para1.add_run("学校")
    p3.font.size = docx.shared.Pt(9)

    # 为了保证年级和班级为空时，不报错，并在该位置添加下划线
    if table.cell_value(i, 6):
        grade_name = table.cell_value(i, 6)
        grade_name_center = text_center_4(str(grade_name))
        p12 = para1.add_run(grade_name_center)
        p12.font.underline = True
        p12.font.size = docx.shared.Pt(9)
    else:
        p14 = para1.add_run("        ")
        p14.font.size = docx.shared.Pt(9)
        p14.font.underline = True

    p13 = para1.add_run("年级")
    p13.font.size = docx.shared.Pt(9)

    #为了保证年级和班级为空时，不报错，并在该位置添加下划线
    class_name = table.cell_value(i, 7)
    if class_name:
        if type(class_name) == float:
            class_name = str(int(class_name))
        else:
            class_name = str(class_name)
        class_center = text_center_4(class_name)
        p4 = para1.add_run(class_center)
        p4.font.underline = True
        p4.font.size = docx.shared.Pt(9)
    else:
        p15 = para1.add_run("        ")
        p15.font.size = docx.shared.Pt(9)
        p15.font.underline = True

    p5 = para1.add_run("班学生姓名 ")
    p5.font.size = docx.shared.Pt(9)

    p6 = para1.add_run(name)
    p6.font.underline = True
    p6.font.size = docx.shared.Pt(9)

    p7 = para1.add_run("，出生日期：")
    p7.font.size = docx.shared.Pt(9)

    p8 = para1.add_run(birthday)
    p8.font.underline = True
    p8.font.size = docx.shared.Pt(9)

    p9 = para1.add_run("，性别：")
    p9.font.size = docx.shared.Pt(9)

    p10 = para1.add_run(sex)
    p10.font.underline = True
    p10.font.size = docx.shared.Pt(9)

    p11 = para1.add_run("，投保本保险壹份。")
    p11.font.size = docx.shared.Pt(9)

    str_output2 = "    监护人（签字）：                是学生的□父亲  □母亲  □其他____________  联系电话:_____________"
    para2 = document.add_paragraph()
    para2.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.SINGLE
    para2_run = para2.add_run(str_output2)
    para2_run.font.size = docx.shared.Pt(9)
    para2_run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)

    str_output3 = "签字日期：2022 年 2 月 28 日"
    para3 = document.add_paragraph()
    para3.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.SINGLE
    para3.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
    para3_run = para3.add_run(str_output3)
    para3_run.font.size = docx.shared.Pt(8)
    para3_run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    document.add_paragraph()
    print(i)
    i += 1

print("总数：", rows)
document.save("demo.doc")
